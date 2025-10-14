# backend/word_parser.py
"""
Word-Dokument Parser für automatische Kapitel- und Szenen-Erkennung
"""
import re
from docx import Document
from typing import List, Dict, Any


def parse_word_document(file_stream) -> Dict[str, Any]:
    """
    Parst ein Word-Dokument und erkennt Kapitel und Szenen automatisch.

    Erkennungslogik:
    - Kapitel: Überschriften (Heading 1, Heading 2) oder Zeilen mit "Kapitel X" / "Chapter X"
    - Szenen: Abschnitte zwischen Kapitel-Überschriften oder durch "***" / "---" getrennt

    Returns:
        {
            "chapters": [
                {
                    "title": "Kapitel 1",
                    "order_index": 0,
                    "scenes": [
                        {
                            "title": "Szene 1",
                            "content": "...",
                            "order_index": 0
                        }
                    ]
                }
            ]
        }
    """
    try:
        doc = Document(file_stream)
    except Exception as e:
        raise ValueError(f"Konnte Word-Dokument nicht lesen: {str(e)}")

    chapters = []
    current_chapter = None
    current_scene_content = []
    scene_counter = 0

    # Patterns für Kapitel-Erkennung
    chapter_patterns = [
        re.compile(r'^(Kapitel|Chapter|Teil|Part)\s+(\d+|[IVXLCDM]+)[\s:]*(.*)$', re.IGNORECASE),
        re.compile(r'^(\d+)\.?\s+(.+)$'),  # "1. Titel" oder "1 Titel"
    ]

    # Scene separator patterns
    scene_separator_pattern = re.compile(r'^\s*[\*\-_]{3,}\s*$')

    def is_chapter_heading(paragraph) -> tuple:
        """Prüft ob Paragraph eine Kapitel-Überschrift ist. Returns (is_chapter, title)"""
        # Check if it's a heading style
        if paragraph.style.name.startswith('Heading 1') or paragraph.style.name.startswith('Heading 2'):
            return True, paragraph.text.strip()

        # Check if it matches chapter patterns
        text = paragraph.text.strip()
        if not text:
            return False, None

        for pattern in chapter_patterns:
            match = pattern.match(text)
            if match:
                # Extrahiere Titel
                if len(match.groups()) >= 3:
                    title = match.group(3).strip() or f"{match.group(1)} {match.group(2)}"
                elif len(match.groups()) >= 2:
                    title = match.group(2).strip()
                else:
                    title = text
                return True, title

        return False, None

    def save_current_scene():
        """Speichert die aktuelle Szene"""
        nonlocal scene_counter, current_scene_content
        if current_chapter is not None and current_scene_content:
            content = '\n\n'.join(current_scene_content).strip()
            if content:  # Nur wenn Inhalt vorhanden
                scene_title = f"Szene {scene_counter + 1}"
                # Versuche ersten Satz als Titel zu verwenden (max 60 Zeichen)
                first_line = current_scene_content[0].strip()
                if first_line and len(first_line) <= 60:
                    scene_title = first_line
                elif first_line:
                    scene_title = first_line[:57] + "..."

                current_chapter["scenes"].append({
                    "title": scene_title,
                    "content": content,
                    "order_index": scene_counter
                })
                scene_counter += 1
        current_scene_content = []

    def start_new_chapter(title: str):
        """Startet ein neues Kapitel"""
        nonlocal current_chapter, scene_counter
        save_current_scene()  # Speichere vorherige Szene

        current_chapter = {
            "title": title,
            "order_index": len(chapters),
            "scenes": []
        }
        chapters.append(current_chapter)
        scene_counter = 0

    # Durchlaufe alle Paragraphen
    for para in doc.paragraphs:
        text = para.text.strip()

        # Skip leere Paragraphen
        if not text:
            continue

        # Prüfe ob Kapitel-Überschrift
        is_chapter, chapter_title = is_chapter_heading(para)
        if is_chapter:
            start_new_chapter(chapter_title)
            continue

        # Prüfe ob Szenen-Separator (***  oder ---)
        if scene_separator_pattern.match(text):
            save_current_scene()
            continue

        # Füge zum aktuellen Szeneninhalt hinzu
        current_scene_content.append(text)

    # Speichere letzte Szene
    save_current_scene()

    # Falls keine Kapitel erkannt wurden, erstelle ein Standard-Kapitel
    if not chapters:
        chapters.append({
            "title": "Kapitel 1",
            "order_index": 0,
            "scenes": []
        })
        current_chapter = chapters[0]

        # Füge allen Content als eine Szene hinzu
        all_content = '\n\n'.join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        if all_content:
            current_chapter["scenes"].append({
                "title": "Szene 1",
                "content": all_content,
                "order_index": 0
            })

    # Entferne leere Kapitel
    chapters = [ch for ch in chapters if ch["scenes"]]

    # Falls immer noch keine Kapitel, erstelle ein Default-Kapitel mit allem Content
    if not chapters:
        all_text = '\n\n'.join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        if all_text:
            chapters.append({
                "title": "Kapitel 1",
                "order_index": 0,
                "scenes": [{
                    "title": "Szene 1",
                    "content": all_text,
                    "order_index": 0
                }]
            })

    return {"chapters": chapters}


def extract_plain_text(file_stream) -> str:
    """Extrahiert nur den reinen Text aus einem Word-Dokument"""
    try:
        doc = Document(file_stream)
        return '\n\n'.join(para.text.strip() for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        raise ValueError(f"Konnte Word-Dokument nicht lesen: {str(e)}")
